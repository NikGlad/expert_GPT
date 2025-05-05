import tkinter as tk
from tkinter import messagebox
from docx import Document

def extract_sections(doc, keywords):
    result = []
    capture = False
    current_keyword = None

    for para in doc.paragraphs:
        text = para.text.strip()

        for keyword in keywords:
            if keyword in text and not text.endswith("конец"):
                capture = True
                current_keyword = keyword
                result.append((para, keyword))
                break
            elif keyword in text and text.endswith("конец") and capture and current_keyword == keyword:
                result.append((para, keyword))
                capture = False
                current_keyword = None
                break
        else:
            if capture:
                result.append((para, current_keyword))

    return result

def generate_doc():
    try:
        src_doc = Document("исходный.docx")
        dst_doc = Document()

    # ПРАВИТЬ
        selected_keywords = []
        if var_vstroyka.get():
            selected_keywords.append("Встройка")
        if var_lift.get():
            selected_keywords.append("Лифт")
        if var_krovlya.get():
            selected_keywords.append("Кровля")
        if var_stoynka.get():
            selected_keywords.append("Автостоянки")

        if not selected_keywords:
            messagebox.showwarning("Внимание", "Выберите хотя бы один раздел.")
            return

        sections = extract_sections(src_doc, selected_keywords)

        for para, _ in sections:
            new_para = dst_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb if run.font.color else None

        dst_doc.save("готовый.docx")
        messagebox.showinfo("Готово", "Файл 'готовый.docx' успешно создан.")
    except Exception as e:
        messagebox.showerror("Ошибка", str(e))


# === Интерфейс ===
root = tk.Tk()
root.title("Expert")
root.geometry("400x400")

# Верхняя часть — прокручиваемая
top_frame = tk.Frame(root)
top_frame.pack(fill="both", expand=True)

canvas = tk.Canvas(top_frame)
scrollbar = tk.Scrollbar(top_frame, orient="vertical", command=canvas.yview)
scrollable_frame = tk.Frame(canvas)

scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)

canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")

# Чекбоксы
# ПРАВИТЬ
var_vstroyka = tk.BooleanVar()
var_lift = tk.BooleanVar()
var_krovlya = tk.BooleanVar()
var_stoynka = tk.BooleanVar()

# ПРАВИТЬ
tk.Label(scrollable_frame, text="Выберите разделы для извлечения:").pack(anchor="w", pady=(5, 5))
tk.Checkbutton(scrollable_frame, text="Встройка", variable=var_vstroyka).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Лифт", variable=var_lift).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Кровля", variable=var_krovlya).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Автостоянки", variable=var_stoynka).pack(anchor="w")


# Нижняя часть — кнопка
bottom_frame = tk.Frame(root)
bottom_frame.pack(fill="x", pady=10)

tk.Button(bottom_frame, text="Создать готовый.docx", command=generate_doc,
          height=2, font=("Arial", 11, "bold")).pack()

# ВАЖНО: запуск интерфейса
root.mainloop()
