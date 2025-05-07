# Импортируем модули для создания интерфейса и работы с Word-документами
import tkinter as tk
from tkinter import messagebox
from docx import Document

# Функция для извлечения нужных разделов из исходного .docx по ключевым словам
def extract_sections(doc, keywords):
    result = []              # Список для хранения найденных параграфов
    capture = False          # Флаг захвата текста
    current_keyword = None   # Текущий активный раздел

    # Перебираем все параграфы документа
    for para in doc.paragraphs:
        text = para.text.strip()  # Удаляем пробелы

        # Проверяем начало и конец каждого раздела по ключевым словам
        for keyword in keywords:
            if keyword in text and not text.endswith("конец"):
                capture = True
                current_keyword = keyword
                result.append((para, keyword))  # Сохраняем параграф
                break
            elif keyword in text and text.endswith("конец") and capture and current_keyword == keyword:
                result.append((para, keyword))  # Добавляем завершающий параграф
                capture = False
                current_keyword = None
                break
        else:
            if capture:
                result.append((para, current_keyword))  # Добавляем текст внутри секции

    return result  # Возвращаем список выбранных параграфов

# Функция создания нового документа из выбранных разделов
def generate_doc():
    try:
        src_doc = Document("исходный.docx")  # Загружаем исходный документ
        dst_doc = Document()                 # Создаем новый пустой документ

        selected_keywords = []              # Список выбранных разделов
        if var_vstroyka.get():
            selected_keywords.append("Встройка")
        if var_lift.get():
            selected_keywords.append("Лифт")
        if var_krovlya.get():
            selected_keywords.append("Кровля")

        if not selected_keywords:
            # Если ничего не выбрано — предупреждение
            messagebox.showwarning("Внимание", "Выберите хотя бы один раздел.")
            return

        # Извлекаем текст из выбранных разделов
        sections = extract_sections(src_doc, selected_keywords)

        # Копируем текст и стили из исходного в новый документ
        for para, _ in sections:
            new_para = dst_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb if run.font.color else None

        dst_doc.save("готовый.docx")  # Сохраняем итоговый файл
        messagebox.showinfo("Готово", "Файл 'готовый.docx' успешно создан.")
    except Exception as e:
        # Показываем сообщение об ошибке, если что-то пошло не так
        messagebox.showerror("Ошибка", str(e))


# === Создание интерфейса ===
root = tk.Tk()                    # Главное окно
root.title("Извлечение текста из Word")  # Заголовок окна
root.geometry("400x400")         # Размер окна

# Верхняя рамка с прокруткой
top_frame = tk.Frame(root)
top_frame.pack(fill="both", expand=True)

canvas = tk.Canvas(top_frame)    # Область прокрутки
scrollbar = tk.Scrollbar(top_frame, orient="vertical", command=canvas.yview)
scrollable_frame = tk.Frame(canvas)

# Обновляем область прокрутки при изменении содержимого
scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)

# Встраиваем прокручиваемый фрейм внутрь канваса
canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)  # Размещение канваса
scrollbar.pack(side="right", fill="y")              # Размещение ползунка

# Переменные для чекбоксов
var_vstroyka = tk.BooleanVar()
var_lift = tk.BooleanVar()
var_krovlya = tk.BooleanVar()

# Подпись и чекбоксы для выбора разделов
tk.Label(scrollable_frame, text="Выберите разделы для извлечения:").pack(anchor="w", pady=(5, 5))
tk.Checkbutton(scrollable_frame, text="Встройка", variable=var_vstroyka).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Лифт", variable=var_lift).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Кровля", variable=var_krovlya).pack(anchor="w")

# Нижняя рамка — кнопка запуска
bottom_frame = tk.Frame(root)
bottom_frame.pack(fill="x", pady=10)

# Кнопка создания нового документа
tk.Button(bottom_frame, text="Создать готовый.docx", command=generate_doc,
          height=2, font=("Arial", 11, "bold")).pack()

# Запуск графического интерфейса
root.mainloop()

# Напоминание:
# Чтобы встроить иконку в .exe, используй:

# pyinstaller --onefile --noconsole --icon=icon.ico main.py
